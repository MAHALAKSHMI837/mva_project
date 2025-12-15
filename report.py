# shortened
# report.py
from docx import Document
from docx.shared import Inches
from pathlib import Path
from utils import setup_logger, ensure_dir
import textwrap
import json

logger = setup_logger("report")
OUT_DIR = ensure_dir("data/reports")

def simple_summary(text, max_words=30):
    words = text.strip().split()
    if len(words) <= max_words:
        return text.strip()
    return " ".join(words[:max_words]) + "..."

def generate_report_docx(video_path, events, segments, out_name=None):
    """Generate comprehensive meeting documentation report"""
    from models.summarizer import extract_key_points
    
    doc = Document()
    doc.add_heading("Meeting Documentation Report", level=1)
    doc.add_paragraph(f"Video: {Path(video_path).name}")
    doc.add_paragraph(f"Scene Changes: {len(events)} | Audio Segments: {len(segments)}")
    doc.add_paragraph(f"Content Type: {'Speech detected' if segments else 'No speech (music/silent video)'}")
    doc.add_paragraph("-" * 50)
    
    if out_name is None:
        out_name = OUT_DIR / (Path(video_path).stem + "_report.docx")
    
    # Key points summary
    key_points = extract_key_points(segments)
    if key_points:
        doc.add_heading("Key Discussion Points", level=2)
        for kp in key_points[:5]:  # Top 5 points
            doc.add_paragraph(f"• {kp['timestamp']:.1f}s: {kp['summary']}")
    else:
        doc.add_heading("Audio Analysis", level=2)
        if not segments:
            doc.add_paragraph("• No speech detected in this video")
            doc.add_paragraph("• Video contains music, background noise, or is silent")
        else:
            doc.add_paragraph("• Audio segments processed but no clear speech found")
    
    # Detailed events
    doc.add_heading("Detailed Timeline", level=2)
    for i, ev in enumerate(events, 1):
        ts, frame = ev["timestamp"], ev["frame"]
        
        # Find transcript
        texts = [seg["text"].strip() for seg in segments 
                if abs(seg["start"] - ts) <= 3]
        caption = " ".join(texts) if texts else "[No speech at this timestamp]"
        
        doc.add_heading(f"{ts:.1f}s", level=3)
        try:
            doc.add_picture(frame, width=Inches(4))
        except:
            doc.add_paragraph(f"[Screenshot: {Path(frame).name}]")
        
        # Always show transcript info
        doc.add_paragraph(f"Transcript: {simple_summary(caption, 40)}")
        doc.add_paragraph(f"Scene Change: Visual content changed at this timestamp")
        
    doc.save(str(out_name))
    logger.info(f"Report saved: {out_name}")
    return str(out_name)

# Optionally we could implement PDF creation using reportlab; docx is typically fine for submission.
